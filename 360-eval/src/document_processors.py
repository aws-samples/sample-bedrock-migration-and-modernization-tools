import os
import json
import logging
import pandas as pd
from typing import List, Dict, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)


# ----------------------------------------
# File Format Detection
# ----------------------------------------

def detect_format(file_path: str) -> str:
    """
    Detect file format based on extension.

    Args:
        file_path: Path to file

    Returns:
        Format string: 'csv', 'json', 'word', 'txt', 'markdown', or 'unknown'
    """
    extension = Path(file_path).suffix.lower()

    format_map = {
        '.csv': 'csv',
        '.json': 'json',
        '.jsonl': 'jsonl',
        '.doc': 'word',
        '.docx': 'word',
        '.txt': 'txt',
        '.md': 'markdown',
        '.markdown': 'markdown',
        '.pdf': 'pdf'
    }

    detected = format_map.get(extension, 'unknown')
    logger.info(f"Detected format '{detected}' for file: {file_path}")
    return detected


# ----------------------------------------
# CSV Parsing
# ----------------------------------------

def parse_csv(file_path: str, text_column: str = None) -> List[str]:
    """
    Parse CSV file and extract text content with structure preservation.

    Args:
        file_path: Path to CSV file
        text_column: Name of column containing text (if specified, only extracts this column)

    Returns:
        List of text strings in key-value format

    Examples:
        >>> # Key-value format (preserves column context)
        >>> parse_csv("regions.csv")
        ['Region Code: us-east-1\\nRegion Name: US East\\nLocation: United States',
         'Region Code: us-west-2\\nRegion Name: US West\\nLocation: United States']
    """
    try:
        logger.info(f"Parsing CSV file: {file_path}")

        # Try standard CSV parsing first
        try:
            df = pd.read_csv(file_path)
        except pd.errors.ParserError as e:
            # Retry with error-tolerant settings
            logger.warning(f"CSV parsing failed with error: {str(e)}")
            logger.warning("Retrying with error-tolerant settings (skipping bad lines)")
            df = pd.read_csv(
                file_path,
                on_bad_lines='skip',  # Skip malformed rows
                encoding_errors='ignore',  # Ignore encoding errors
                dtype=str  # Read everything as string to avoid type issues
            )
        except UnicodeDecodeError:
            # Try different encoding
            logger.warning("UTF-8 decoding failed, trying latin-1 encoding")
            df = pd.read_csv(
                file_path,
                encoding='latin-1',
                on_bad_lines='skip',
                dtype=str
            )

        if df.empty:
            logger.warning(f"CSV file is empty or all rows were skipped: {file_path}")
            return []

        # Handle single column extraction
        if text_column:
            if text_column not in df.columns:
                raise ValueError(f"Column '{text_column}' not found in CSV. Available: {list(df.columns)}")
            texts = df[text_column].dropna().astype(str).tolist()
            logger.info(f"Extracted {len(texts)} text entries from column '{text_column}'")
            return texts

        # Handle multi-column formatting (key-value format)
        texts = []

        # Key-value format: preserves column context
        for idx, row in df.iterrows():
            row_lines = []
            for col_name, value in row.items():
                if pd.notna(value):  # Skip NaN values
                    # Clean up value
                    value_str = str(value).strip()
                    if value_str:  # Skip empty strings
                        row_lines.append(f"{col_name}: {value_str}")

            if row_lines:  # Only add non-empty rows
                row_text = "\n".join(row_lines)
                texts.append(row_text)

        logger.info(f"Extracted {len(texts)} text entries from CSV using key-value format")
        return texts

    except Exception as e:
        logger.error(f"Error parsing CSV file: {str(e)}", exc_info=True)
        raise


# ----------------------------------------
# JSON Parsing
# ----------------------------------------

def parse_json(file_path: str, text_field: str = None) -> List[str]:
    """
    Parse JSON file and extract text content.

    Args:
        file_path: Path to JSON file
        text_field: Field name containing text (supports nested fields with dot notation)

    Returns:
        List of text strings
    """
    try:
        logger.info(f"Parsing JSON file: {file_path}")
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        texts = []

        # If data is a list of objects
        if isinstance(data, list):
            for item in data:
                text = extract_text_from_json(item, text_field)
                if text:
                    texts.append(text)
        # If data is a single object
        elif isinstance(data, dict):
            text = extract_text_from_json(data, text_field)
            if text:
                texts.append(text)
        else:
            texts.append(str(data))

        logger.info(f"Extracted {len(texts)} text entries from JSON")
        return texts

    except Exception as e:
        logger.error(f"Error parsing JSON file: {str(e)}", exc_info=True)
        raise


def extract_text_from_json(obj: Dict, field: str = None) -> str:
    """
    Extract text from JSON object.

    Args:
        obj: JSON object (dict)
        field: Field name to extract (supports dot notation like 'data.content')

    Returns:
        Extracted text string
    """
    if field:
        # Support nested fields with dot notation
        keys = field.split('.')
        value = obj
        for key in keys:
            if isinstance(value, dict) and key in value:
                value = value[key]
            else:
                return ""
        return str(value) if value else ""
    else:
        # Concatenate all string values in the object
        def extract_strings(data):
            if isinstance(data, str):
                return [data]
            elif isinstance(data, dict):
                strings = []
                for v in data.values():
                    strings.extend(extract_strings(v))
                return strings
            elif isinstance(data, list):
                strings = []
                for item in data:
                    strings.extend(extract_strings(item))
                return strings
            else:
                return [str(data)] if data else []

        strings = extract_strings(obj)
        return ' '.join(strings)


def parse_jsonl(file_path: str, text_field: str = None) -> List[str]:
    """
    Parse JSONL file (one JSON object per line).

    Args:
        file_path: Path to JSONL file
        text_field: Field name containing text

    Returns:
        List of text strings
    """
    try:
        logger.info(f"Parsing JSONL file: {file_path}")
        texts = []

        with open(file_path, 'r', encoding='utf-8') as f:
            for line_num, line in enumerate(f, 1):
                if line.strip():
                    try:
                        obj = json.loads(line)
                        text = extract_text_from_json(obj, text_field)
                        if text:
                            texts.append(text)
                    except json.JSONDecodeError as e:
                        logger.warning(f"Skipping invalid JSON at line {line_num}: {str(e)}")

        logger.info(f"Extracted {len(texts)} text entries from JSONL")
        return texts

    except Exception as e:
        logger.error(f"Error parsing JSONL file: {str(e)}", exc_info=True)
        raise


# ----------------------------------------
# Word Document Parsing
# ----------------------------------------

def parse_word(file_path: str) -> List[str]:
    """
    Parse Word document (.doc/.docx) and extract text.

    Args:
        file_path: Path to Word document

    Returns:
        List of text strings (one per paragraph)
    """
    try:
        from docx import Document

        logger.info(f"Parsing Word document: {file_path}")
        doc = Document(file_path)

        texts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)

        # Also extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join(cell.text.strip() for cell in row.cells)
                if row_text:
                    texts.append(row_text)
                    table_texts.append(row_text)

        logger.info(f"Extracted {len(texts)} text entries from Word document")
        return texts

    except ImportError:
        logger.error("python-docx library not installed. Install with: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"Error parsing Word document: {str(e)}", exc_info=True)
        raise


# ----------------------------------------
# Text File Parsing
# ----------------------------------------

def parse_txt(file_path: str) -> List[str]:
    """
    Parse plain text file.

    Args:
        file_path: Path to text file

    Returns:
        List of text strings (one per non-empty line)
    """
    try:
        logger.info(f"Parsing text file: {file_path}")
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = [line.strip() for line in f if line.strip()]

        logger.info(f"Extracted {len(lines)} text entries from text file")
        return lines

    except Exception as e:
        logger.error(f"Error parsing text file: {str(e)}", exc_info=True)
        raise


# ----------------------------------------
# Markdown Parsing
# ----------------------------------------

def parse_markdown(file_path: str) -> List[str]:
    """
    Parse Markdown file and extract text.

    Args:
        file_path: Path to Markdown file

    Returns:
        List of text strings (one per paragraph/section)
    """
    try:
        import markdown
        from markdown.extensions import Extension
        from markdown.treeprocessors import Treeprocessor

        logger.info(f"Parsing Markdown file: {file_path}")

        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Convert Markdown to HTML
        md = markdown.Markdown()
        html = md.convert(md_content)

        # Extract text from HTML
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.texts = []
                self.current_text = []

            def handle_data(self, data):
                text = data.strip()
                if text:
                    self.current_text.append(text)

            def handle_endtag(self, tag):
                if tag in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li']:
                    if self.current_text:
                        self.texts.append(' '.join(self.current_text))
                        self.current_text = []

        parser = TextExtractor()
        parser.feed(html)

        # Add any remaining text
        if parser.current_text:
            parser.texts.append(' '.join(parser.current_text))

        # If no structured parsing worked, fall back to line-based
        if not parser.texts:
            parser.texts = [line.strip() for line in md_content.split('\n') if line.strip() and not line.strip().startswith('#')]

        logger.info(f"Extracted {len(parser.texts)} text entries from Markdown")
        return parser.texts

    except ImportError:
        logger.warning("markdown library not installed. Falling back to plain text parsing.")
        return parse_txt(file_path)
    except Exception as e:
        logger.error(f"Error parsing Markdown file: {str(e)}", exc_info=True)
        raise


# ----------------------------------------
# PDF Parsing
# ----------------------------------------

def parse_pdf(
    file_path: str
) -> List[str]:
    """
    Parse PDF file and extract text.

    Args:
        file_path: Path to PDF file

    Returns:
        List of text strings (one per page)

    Example:
        texts = parse_pdf('doc.pdf')
    """
    try:
        from PyPDF2 import PdfReader

        logger.info(f"Parsing PDF file: {file_path}")
        reader = PdfReader(file_path)

        texts = []

        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text().strip()

            if text:
                texts.append(text)

        logger.info(f"Extracted {len(texts)} text pages from PDF ({len(reader.pages)} pages)")
        return texts

    except ImportError:
        logger.error("PyPDF2 library not installed. Install with: pip install PyPDF2")
        raise
    except Exception as e:
        logger.error(f"Error parsing PDF file: {str(e)}", exc_info=True)
        raise


# ----------------------------------------
# Universal Document Parser
# ----------------------------------------

def parse_document(
    file_path: str,
    file_format: str = None,
    text_field: str = None
) -> List[str]:
    """
    Parse document of any supported format, optionally with image descriptions.

    Args:
        file_path: Path to document
        file_format: Format override (if None, auto-detects)
        text_field: Field/column name for structured formats

    Returns:
        List of text strings
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    # Auto-detect format if not specified or set to 'auto'
    if not file_format or file_format == 'auto':
        file_format = detect_format(file_path)

    logger.info(f"Parsing document as {file_format}: {file_path}")

    # Route to appropriate parser
    if file_format == 'csv':
        return parse_csv(file_path, text_field)
    elif file_format == 'json':
        return parse_json(file_path, text_field)
    elif file_format == 'jsonl':
        return parse_jsonl(file_path, text_field)
    elif file_format == 'word':
        return parse_word(file_path)
    elif file_format == 'txt':
        return parse_txt(file_path)
    elif file_format == 'markdown':
        return parse_markdown(file_path)
    elif file_format == 'pdf':
        return parse_pdf(file_path)
    else:
        logger.warning(f"Unknown format '{file_format}', treating as plain text")
        return parse_txt(file_path)


# ----------------------------------------
# Bulk Processing
# ----------------------------------------

def parse_multiple_documents(
    file_paths: List[str],
    file_format: str = None,
    text_field: str = None
) -> List[Dict[str, any]]:
    """
    Parse multiple documents and track source.

    Args:
        file_paths: List of file paths to parse
        file_format: Format override for all files
        text_field: Field/column name for structured formats

    Returns:
        List of dicts with 'text' and 'source' keys
    """
    results = []

    for file_path in file_paths:
        try:
            texts = parse_document(file_path, file_format, text_field)
            for text in texts:
                results.append({
                    'text': text,
                    'source': file_path,
                    'source_file': os.path.basename(file_path)
                })
        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {str(e)}")
            continue

    logger.info(f"Parsed {len(file_paths)} files, extracted {len(results)} text entries")
    return results
